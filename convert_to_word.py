#!/usr/bin/env python3
"""
Convert the Detailed Project Summary Report to Word Document Format
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def create_word_document():
    # Create a new Document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title Page
    title = doc.add_heading('🚨 DETAILED PROJECT SUMMARY REPORT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Educational Cybersecurity Platform - Comprehensive Analysis', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add page break after title
    doc.add_page_break()
    
    # Read the markdown file content
    with open('/app/DETAILED_PROJECT_SUMMARY_REPORT.md', 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Process the content
    lines = content.split('\n')
    current_section = None
    
    for line in lines:
        line = line.strip()
        
        if not line:
            continue
            
        # Skip the title lines we already added
        if line.startswith('# 🚨 DETAILED PROJECT SUMMARY REPORT') or \
           line.startswith('## Educational Cybersecurity Platform - Comprehensive Analysis'):
            continue
            
        # Main sections (## 1.0, ## 2.0, etc.)
        if re.match(r'^## \d+\.\d+ [A-Z]', line):
            section_text = line.replace('## ', '').replace('**', '')
            doc.add_heading(section_text, level=1)
            
        # Sub-sections (### )
        elif line.startswith('### '):
            subsection_text = line.replace('### ', '').replace('**', '')
            doc.add_heading(subsection_text, level=2)
            
        # Sub-sub-sections (#### )
        elif line.startswith('#### '):
            subsubsection_text = line.replace('#### ', '').replace('**', '')
            doc.add_heading(subsubsection_text, level=3)
            
        # Code blocks
        elif line.startswith('```'):
            continue  # Skip code block markers
            
        # Bullet points with various markers
        elif line.startswith('- ') or line.startswith('* ') or line.startswith('+ '):
            bullet_text = line[2:].replace('**', '').replace('✅ ', '✅ ').replace('❌ ', '❌ ')
            p = doc.add_paragraph(bullet_text, style='List Bullet')
            
        # Numbered lists
        elif re.match(r'^\d+\. ', line):
            numbered_text = re.sub(r'^\d+\. ', '', line).replace('**', '')
            p = doc.add_paragraph(numbered_text, style='List Number')
            
        # Tree structure (├── └── │)
        elif '├──' in line or '└──' in line or '│' in line:
            tree_text = line.replace('**', '')
            p = doc.add_paragraph(tree_text)
            # Set monospace font for tree structures
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                
        # Horizontal rules
        elif line.startswith('---'):
            doc.add_paragraph().add_run().add_break()
            
        # Regular paragraphs
        elif line and not line.startswith('#'):
            # Clean up markdown formatting
            clean_text = line.replace('**', '').replace('*', '').replace('`', '')
            
            # Skip certain formatting lines
            if clean_text.startswith('```') or clean_text == '---':
                continue
                
            # Add as paragraph if it has substantial content
            if len(clean_text.strip()) > 1:
                p = doc.add_paragraph(clean_text)
    
    # Add final section with metadata
    doc.add_page_break()
    
    final_section = doc.add_heading('📊 DOCUMENT METADATA', level=1)
    
    metadata_content = [
        "Document Type: Comprehensive Project Analysis Report",
        "Platform: Educational Cybersecurity Learning Environment",
        "Technology Stack: FastAPI + React + MongoDB",
        "Security Features: AES Encryption, Multi-client Architecture, Real-time Monitoring",
        "Testing Results: 16/18 tests passed (89% success rate)",
        "Documentation Suite: 7 comprehensive guides and references",
        "Educational Framework: Complete responsible use and safety guidelines",
        "Target Audience: Students, researchers, cybersecurity professionals",
        "Usage: Authorized educational and research purposes only",
        "",
        "⚠️ IMPORTANT: This document describes an educational cybersecurity platform designed for authorized learning and research purposes only. All users must comply with applicable laws and regulations."
    ]
    
    for item in metadata_content:
        if item:
            doc.add_paragraph(f"• {item}")
        else:
            doc.add_paragraph("")
    
    # Save the document
    doc.save('/app/Detailed_Project_Summary_Report.docx')
    print("✅ Word document created successfully: /app/Detailed_Project_Summary_Report.docx")

if __name__ == "__main__":
    create_word_document()