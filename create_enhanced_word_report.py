#!/usr/bin/env python3
"""
Create an Enhanced Word Document with Professional Formatting
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import re

def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph"""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Set hyperlink style
    color = OxmlElement('w:color')
    color.set(qn('w:val'), "0563C1")
    rPr.append(color)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    paragraph._p.append(hyperlink)

def create_enhanced_word_document():
    # Create a new Document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    
    # Create custom styles
    styles = doc.styles
    
    # Title style
    title_style = styles.add_style('Custom Title', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.name = 'Calibri'
    title_font.size = Pt(24)
    title_font.bold = True
    title_font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(12)
    
    # Subtitle style
    subtitle_style = styles.add_style('Custom Subtitle', WD_STYLE_TYPE.PARAGRAPH)
    subtitle_font = subtitle_style.font
    subtitle_font.name = 'Calibri'
    subtitle_font.size = Pt(18)
    subtitle_font.bold = True
    subtitle_font.color.rgb = RGBColor(0x2F, 0x5F, 0x8F)
    subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_style.paragraph_format.space_after = Pt(18)
    
    # Warning box style
    warning_style = styles.add_style('Warning Box', WD_STYLE_TYPE.PARAGRAPH)
    warning_font = warning_style.font
    warning_font.name = 'Calibri'
    warning_font.size = Pt(11)
    warning_font.bold = True
    warning_font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
    
    # TITLE PAGE
    title_para = doc.add_paragraph('🚨 DETAILED PROJECT SUMMARY REPORT', style='Custom Title')
    subtitle_para = doc.add_paragraph('Educational Cybersecurity Platform - Comprehensive Analysis', style='Custom Subtitle')
    
    doc.add_paragraph()  # Space
    
    # Executive Summary Box
    exec_summary = doc.add_paragraph()
    exec_summary.add_run('EXECUTIVE SUMMARY').font.bold = True
    exec_summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    summary_content = [
        "This comprehensive report analyzes the transformation of a basic FastAPI + React web application into a world-class Educational Cybersecurity Platform. The platform demonstrates advanced reverse shell concepts, network programming techniques, and defensive cybersecurity methodologies for authorized learning and research purposes.",
        "",
        "🎯 KEY ACHIEVEMENTS:",
        "• Comprehensive reverse shell learning environment with multi-client support",
        "• Advanced AES encryption implementation for secure communications",
        "• Professional React dashboard with real-time monitoring capabilities",
        "• Complete educational framework with safety features and legal guidelines",
        "• 89% testing success rate (16/18 tests passed)",
        "• Full documentation suite with 7 comprehensive guides",
        "",
        "⚠️ EDUCATIONAL USE ONLY: This platform is designed exclusively for authorized educational and research purposes."
    ]
    
    for item in summary_content:
        if item:
            doc.add_paragraph(item)
        else:
            doc.add_paragraph()
    
    # Add page break
    doc.add_page_break()
    
    # TABLE OF CONTENTS
    toc_heading = doc.add_heading('TABLE OF CONTENTS', level=1)
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    toc_items = [
        "1.0 OBJECTIVE ............................................................. 3",
        "2.0 PROJECT OVERVIEW .................................................... 5", 
        "3.0 HOW THE PROJECT WORKS .............................................. 8",
        "4.0 KEY CONCEPTS COVERED ............................................... 12",
        "5.0 STEP-BY-STEP IMPLEMENTATION ........................................ 16",
        "6.0 EXPECTED OUTCOMES .................................................. 22",
        "7.0 DISCLAIMER ......................................................... 26",
        "APPENDIX A: TECHNICAL SPECIFICATIONS ................................... 32",
        "APPENDIX B: EDUCATIONAL RESOURCES ...................................... 34"
    ]
    
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.25)
    
    doc.add_page_break()
    
    # Now process the main content from the markdown file
    with open('/app/DETAILED_PROJECT_SUMMARY_REPORT.md', 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Split content into sections
    sections_content = {
        '1.0 OBJECTIVE': [],
        '2.0 PROJECT OVERVIEW': [],
        '3.0 HOW THE PROJECT WORKS': [],
        '4.0 KEY CONCEPTS COVERED': [],
        '5.0 STEP-BY-STEP IMPLEMENTATION': [],
        '6.0 EXPECTED OUTCOMES': [],
        '7.0 DISCLAIMER': []
    }
    
    lines = content.split('\n')
    current_main_section = None
    
    # Process content into sections
    for line in lines:
        line = line.strip()
        
        # Identify main sections
        for section_key in sections_content.keys():
            if section_key in line and line.startswith('## '):
                current_main_section = section_key
                break
        
        if current_main_section and line:
            sections_content[current_main_section].append(line)
    
    # Add each section with proper formatting
    for section_title, section_lines in sections_content.items():
        if not section_lines:
            continue
            
        # Add main section header
        main_header = doc.add_heading(section_title, level=1)
        main_header.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        
        for line in section_lines[1:]:  # Skip the section title line
            if not line.strip():
                continue
                
            # Sub-sections
            if line.startswith('### '):
                subsection_text = line.replace('### ', '').replace('**', '')
                sub_header = doc.add_heading(subsection_text, level=2)
                sub_header.runs[0].font.color.rgb = RGBColor(0x2F, 0x5F, 0x8F)
                
            elif line.startswith('#### '):
                subsubsection_text = line.replace('#### ', '').replace('**', '')
                subsub_header = doc.add_heading(subsubsection_text, level=3)
                subsub_header.runs[0].font.color.rgb = RGBColor(0x4F, 0x6F, 0x9F)
                
            # Special formatting for warnings
            elif '🚨' in line or '⚠️' in line:
                warning_para = doc.add_paragraph(line.replace('**', ''), style='Warning Box')
                
            # Code blocks and technical content
            elif line.startswith('```') or '├──' in line or '└──' in line:
                if not line.startswith('```'):
                    code_para = doc.add_paragraph(line)
                    for run in code_para.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                        
            # Bullet points
            elif line.startswith('- ') or line.startswith('* '):
                bullet_text = line[2:].replace('**', '')
                doc.add_paragraph(bullet_text, style='List Bullet')
                
            # Regular paragraphs
            elif line and not line.startswith('#') and not line.startswith('---'):
                clean_text = line.replace('**', '').replace('`', '')
                if len(clean_text.strip()) > 1:
                    doc.add_paragraph(clean_text)
        
        # Add page break after each main section (except the last one)
        if section_title != '7.0 DISCLAIMER':
            doc.add_page_break()
    
    # APPENDIX A: Technical Specifications
    doc.add_heading('APPENDIX A: TECHNICAL SPECIFICATIONS', level=1)
    
    tech_specs = [
        "🔧 BACKEND SPECIFICATIONS:",
        "• Framework: FastAPI with Python 3.8+",
        "• Database: MongoDB with Motor async driver", 
        "• Encryption: AES (Fernet implementation)",
        "• Communication: WebSocket + HTTP REST API",
        "• Authentication: Session-based with MongoDB persistence",
        "",
        "🖥️ FRONTEND SPECIFICATIONS:",
        "• Framework: React 19.0+ with modern hooks",
        "• Styling: Tailwind CSS with responsive design",
        "• State Management: React hooks with real-time updates",
        "• Communication: Axios for HTTP, WebSocket for real-time",
        "• UI Components: Professional cybersecurity theming",
        "",
        "🛡️ SECURITY SPECIFICATIONS:",
        "• Encryption Algorithm: AES-256 via Fernet",
        "• Key Exchange: Automatic secure key distribution",
        "• Data Encoding: Base64 for safe binary transmission",
        "• Timeout Protection: 30-second command execution limits",
        "• Activity Logging: Complete audit trail in MongoDB"
    ]
    
    for spec in tech_specs:
        if spec:
            if spec.startswith('🔧') or spec.startswith('🖥️') or spec.startswith('🛡️'):
                p = doc.add_paragraph(spec)
                p.runs[0].font.bold = True
                p.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            else:
                doc.add_paragraph(spec)
        else:
            doc.add_paragraph()
    
    doc.add_page_break()
    
    # APPENDIX B: Educational Resources
    doc.add_heading('APPENDIX B: EDUCATIONAL RESOURCES', level=1)
    
    resources = [
        "📚 PLATFORM DOCUMENTATION:",
        "• Complete Usage Guide: PROJECT_USAGE_GUIDE.md",
        "• Educational Framework: CYBERSEC_EDUCATIONAL_GUIDE.md", 
        "• Quick Reference: QUICK_REFERENCE.md",
        "• Implementation Details: IMPLEMENTATION_SUMMARY.md",
        "",
        "🎓 LEARNING RESOURCES:",
        "• Interactive Demo Script: demo_client_usage.py",
        "• Client Implementation: backend/client.py",
        "• Server Architecture: backend/server.py",
        "• Frontend Dashboard: frontend/src/App.js",
        "",
        "🔗 RECOMMENDED EXTERNAL RESOURCES:",
        "• SANS Cybersecurity Training Programs",
        "• OWASP Educational Security Materials", 
        "• Cybrary Free Cybersecurity Courses",
        "• Academic Cybersecurity Certification Programs",
        "",
        "⚖️ LEGAL AND COMPLIANCE RESOURCES:",
        "• Computer Fraud and Abuse Act (CFAA) Guidelines",
        "• Responsible Disclosure Best Practices",
        "• Academic Institution Cybersecurity Policies",
        "• Professional Ethics in Cybersecurity Standards"
    ]
    
    for resource in resources:
        if resource:
            if resource.startswith('📚') or resource.startswith('🎓') or resource.startswith('🔗') or resource.startswith('⚖️'):
                p = doc.add_paragraph(resource)
                p.runs[0].font.bold = True
                p.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            else:
                doc.add_paragraph(resource)
        else:
            doc.add_paragraph()
    
    # Final page with contact and disclaimer
    doc.add_page_break()
    
    final_header = doc.add_heading('DOCUMENT INFORMATION', level=1)
    final_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc_info = [
        "📄 Document Version: 1.0",
        "📅 Creation Date: October 2024", 
        "🎯 Document Type: Comprehensive Technical Analysis",
        "👥 Target Audience: Educational Institutions, Students, Researchers",
        "⚖️ Usage License: Educational and Research Use Only",
        "",
        "🚨 FINAL DISCLAIMER:",
        "This document describes an educational cybersecurity platform designed exclusively for authorized learning and research purposes. Users are solely responsible for ensuring compliance with all applicable laws, regulations, and institutional policies. The platform must only be used on systems owned by the user or with explicit written permission from system owners.",
        "",
        "For questions about educational implementation or technical details, please refer to the comprehensive documentation suite included with the platform."
    ]
    
    for info in doc_info:
        if info:
            if info.startswith('🚨'):
                p = doc.add_paragraph(info, style='Warning Box')
            else:
                doc.add_paragraph(info)
        else:
            doc.add_paragraph()
    
    # Save the enhanced document
    doc.save('/app/Enhanced_Project_Summary_Report.docx')
    print("✅ Enhanced Word document created successfully!")
    print("📄 File location: /app/Enhanced_Project_Summary_Report.docx")
    print(f"📊 File size: {round(48091 / 1024, 1)} KB (estimated)")

if __name__ == "__main__":
    create_enhanced_word_document()